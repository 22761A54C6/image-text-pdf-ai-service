import csv
from fastapi import HTTPException

from app.ocr_service import preprocess_image, run_ocr

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


IMAGE_TYPES = {"image/jpeg", "image/png", "image/jpg", "image/webp", "image/bmp"}
PDF_TYPES = {"application/pdf"}
TEXT_TYPES = {"text/plain"}
CSV_TYPES = {"text/csv", "application/vnd.ms-excel"}
DOCX_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


def load_document_text(file_path: str, content_type: str) -> str:
    """Single entry point: dispatches to the right loader by content type,
    returns raw text ready for extract_products_with_llm()."""
    if content_type in IMAGE_TYPES:
        processed = preprocess_image(file_path)
        return run_ocr(processed)
    if content_type in PDF_TYPES:
        return _load_pdf(file_path)
    if content_type in TEXT_TYPES:
        return _load_txt(file_path)
    if content_type in CSV_TYPES:
        return _load_csv(file_path)
    if content_type in DOCX_TYPES:
        return _load_docx(file_path)
    raise HTTPException(status_code=400, detail=f"Unsupported document type: {content_type}")


def _load_pdf(file_path: str) -> str:
    if pdfplumber is None:
        raise HTTPException(status_code=500, detail="pdfplumber not installed")

    lines = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)

    combined = "\n".join(lines).strip()
    if not combined:
        # PDF has no text layer — it's likely a scanned image inside a PDF.
        # Flagging rather than guessing, same principle as your menu-price case earlier.
        raise HTTPException(
            status_code=422,
            detail="No extractable text in PDF — it may be scanned/image-only. "
                   "OCR fallback for scanned PDFs isn't implemented yet."
        )
    return combined


def _load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _load_csv(file_path: str) -> str:
    lines = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        for row in csv.reader(f):
            cells = [c.strip() for c in row if c.strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def _load_docx(file_path: str) -> str:
    if docx is None:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    document = docx.Document(file_path)
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)